import ipaddress
import os
import re
import io
import json
import time
import hashlib
from typing import List, Tuple, Optional, Dict
from urllib.parse import urljoin, urlparse
from typing import Any
import streamlit as st
from tenacity import retry, stop_after_attempt, wait_exponential, RetryError
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import Pt
import requests

ACCESS_LIMIT = 4
ACCESS_WINDOW_SECONDS = 24 * 3600
ALLOWED_CLIENT_IPS: List[str] = [
    
]
# Try to reuse your working llm_client.py (optional)

# Use Gemini API
from gemini_client import call_gemini
# --- Access limit config ---
ACCESS_LIMIT = 4
ACCESS_WINDOW_SECONDS = 24 * 3600
ACCESS_TRACK_FILE = "access_tracker.json"


try:
    from readability import Document  # readability-lxml
    _HAS_READABILITY = True
except Exception:
    Document = None
    _HAS_READABILITY = False

load_dotenv()

def get_user_ip():
    ip = os.environ.get("REMOTE_ADDR")
    if not ip:
        ip = hashlib.md5(str(st.session_state.get('run_id', time.time())).encode()).hexdigest()[:8]
    return ip

def get_client_ip_reliable() -> str:
    """
    Attempts to get the client's IP address by prioritizing X-Forwarded-For header.
    """
    try:
        # 1. Check Streamlit context headers for X-Forwarded-For (most common in deployed Streamlit)
        # This is the most likely location for the true client IP behind a reverse proxy.
        if 'context' in st.session_state and 'headers' in st.session_state.context:
            headers = st.session_state.context.get('headers', {})
            # X-Forwarded-For lists the client IP first, followed by proxies
            x_forwarded_for = headers.get('X-Forwarded-For', '')
            if x_forwarded_for:
                return x_forwarded_for.split(',')[0].strip()

        # 2. Fallback to native Streamlit context (v1.45.0+)
        if hasattr(st, 'context') and hasattr(st.context, 'ip_address'):
             if st.context.ip_address:
                 return st.context.ip_address

    except Exception:
        pass # Fall through to the most basic check

    # 3. Fallback to environment variable
    ip_env = os.environ.get("REMOTE_ADDR")
    if ip_env:
        return ip_env

    # Unreliable fallback
    return "127.0.0.1"


def check_ip_whitelist(client_ip: str):
    """Checks if the client_ip is in the ALLOWED_CLIENT_IPS list or range and stops the app if not."""
    if not ALLOWED_CLIENT_IPS:
        # No IP restriction configured, allow access
        return

    try:
        ip_addr = ipaddress.ip_address(client_ip)
        is_allowed = False

        for allowed in ALLOWED_CLIENT_IPS:
            allowed = allowed.strip()
            if not allowed: continue

            if "/" in allowed:
                # Check against an IP range (CIDR block)
                if ip_addr in ipaddress.ip_network(allowed, strict=False):
                    is_allowed = True
                    break
            elif str(ip_addr) == allowed:
                # Check against a single IP
                is_allowed = True
                break

        if not is_allowed:
            st.error(f"⚠️ Access Restricted: Your IP address ({client_ip}) is not authorized to use this application.")
            st.stop() # Immediately stop the script execution

    except ValueError:
        # Handle cases where the retrieved client_ip is not a valid IP address format
        st.error("⚠️ Access Restricted: Security check failed (invalid IP format).")
        st.stop()

def load_access_tracker():
    try:
        with open(ACCESS_TRACK_FILE, "r") as f:
            return _json.load(f)
    except Exception:
        return {}

def save_access_tracker(tracker):
    with open(ACCESS_TRACK_FILE, "w") as f:
        _json.dump(tracker, f)

def check_and_update_access(ip):
    tracker = load_access_tracker()
    now = int(time.time())
    accesses = tracker.get(ip, [])
    accesses = [t for t in accesses if now - t < ACCESS_WINDOW_SECONDS]
    if len(accesses) >= ACCESS_LIMIT:
        return False, ACCESS_LIMIT, min(accesses) + ACCESS_WINDOW_SECONDS - now
    accesses.append(now)
    tracker[ip] = accesses
    save_access_tracker(tracker)
    return True, ACCESS_LIMIT - len(accesses), None


class ExecSummary(BaseModel):
    url: str
    title: Optional[str]
    product: Optional[str]
    executive_summary: str
    key_highlights: List[str]
    resources_note: Optional[str]
    sample_resources: List[str] = []


HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; ExecSummarizer/1.0; +https://internal.local)"
}

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=8))
def fetch_url(url: str, timeout: int = 20, verify=True) -> Tuple[str, str]:

    resp = requests.get(
        url,
        headers=HEADERS,
        timeout=timeout,
        allow_redirects=True,
        verify=verify,
    )
    resp.raise_for_status()
    # Cap size ~10MB
    if int(resp.headers.get("Content-Length") or 0) > 10_000_000:
        raise RuntimeError("Page too large for summarization.")
    return (resp.url, resp.text)


from bs4 import BeautifulSoup, Tag

def make_soup(html: str) -> BeautifulSoup:
    for parser in ("lxml", "html.parser", "html5lib"):
        try:
            return BeautifulSoup(html or "", parser)
        except Exception:
            continue
    return BeautifulSoup(html or "", "html.parser")


# --- Hardened boilerplate stripper (fixes the original AttributeError) ---
def strip_boilerplate(soup: BeautifulSoup) -> None:
    for tag in soup(["script", "style", "noscript", "template", "iframe"]):
        try: tag.decompose()
        except Exception: pass

    for tag in soup.find_all(["header", "footer", "nav", "aside"]):
        try: tag.decompose()
        except Exception: pass

    patterns = ["breadcrumb", "cookie", "consent", "sidebar", "newsletter",
                "promo", "social", "banner", "advert"]

    for attr in ["class", "id", "role", "aria-label"]:
        for node in list(soup.find_all(True, attrs={attr: True})):
            if not isinstance(node, Tag):
                continue
            try:
                val_attr = node.attrs.get(attr)
                if isinstance(val_attr, (list, tuple)):
                    val = " ".join([str(v) for v in val_attr if v])
                elif isinstance(val_attr, (bool, type(None))):
                    val = ""
                else:
                    val = str(val_attr)
                if val and any(p in val.lower() for p in patterns):
                    node.decompose()
            except Exception:
                continue


# --- Always return a (title, h1) pair ---
def get_title_and_h1(html: str) -> Tuple[Optional[str], Optional[str]]:
    soup = make_soup(html)
    title = soup.title.get_text(strip=True) if soup.title else None
    h1 = soup.find("h1")
    h1_text = h1.get_text(strip=True) if h1 else None
    # Always return a 2‑tuple
    return title, h1_text


# --- Always return (links, pdfs) ---
def collect_links(html: str, base_url: str) -> Tuple[List[str], List[str]]:
    soup = make_soup(html)
    links, pdfs = [], []
    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").strip()
        if not href or href.startswith("#"):
            continue
        abs_url = urljoin(base_url, href)
        (pdfs if abs_url.lower().endswith(".pdf") else links).append(abs_url)

    # dedupe preserve order
    def _dedupe(seq):
        seen, out = set(), []
        for x in seq:
            if x not in seen:
                seen.add(x); out.append(x)
        return out

    return _dedupe(links), _dedupe(pdfs)


# --- Extractors always return strings (never None) ---
def readability_extract(html: str) -> str:
    if not _HAS_READABILITY:
        return simple_extract(html)
    try:
        doc = Document(html or "")
        summary_html = doc.summary(html_partial=True) or ""
        soup = make_soup(summary_html)
        strip_boilerplate(soup)
        text = soup.get_text(separator="\n")
        return clean_text(text)
    except Exception:
        # Guaranteed string fallback
        return simple_extract(html)


def simple_extract(html: str) -> str:
    soup = make_soup(html or "")
    strip_boilerplate(soup)
    text = soup.get_text(separator="\n")
    return clean_text(text)


def clean_text(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()

def readability_extract(html: str) -> str:
    if not _HAS_READABILITY:
        return simple_extract(html)
    doc = Document(html)
    summary_html = doc.summary(html_partial=True)
    soup = make_soup(summary_html)
    strip_boilerplate(soup)
    text = soup.get_text(separator="\n")
    return clean_text(text)

def simple_extract(html: str) -> str:
    soup = make_soup(html)
    strip_boilerplate(soup)
    text = soup.get_text(separator="\n")
    return clean_text(text)


def _ensure_pair(value, label: str):
    """
    Ensure a (x, y) tuple. Raise a clear error if not.
    """
    if isinstance(value, tuple) and len(value) == 2:
        return value
    raise RuntimeError(f"{label} returned invalid value (expected 2‑tuple), got: {type(value).__name__} -> {repr(value)[:200]}")

def infer_product_name(title: Optional[str], h1: Optional[str]) -> Optional[str]:
    use = h1 or title
    if not use:
        return None
    use = re.sub(r"\s*[\-\–—]\s*.*$", "", use).strip()
    if len(use) > 80:
        use = use[:77] + "..."
    return use

def truncate_for_model(text: str, max_chars: int = 12000) -> str:
    if len(text) <= max_chars:
        return text
    head = text[: max_chars // 2]
    tail = text[-max_chars // 2 :]
    return head + "\n...\n" + tail




def _count_words(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text or ""))

def _safe_json_loads(s: str) -> Any:
    try:
        return _json.loads(s)
    except Exception:
        return None

def _extract_ld_json_types(soup: BeautifulSoup) -> List[str]:
    types: List[str] = []
    for tag in soup.find_all("script", type=lambda t: t and "ld+json" in t):
        data = _safe_json_loads(tag.string or "")
        if isinstance(data, dict) and "@type" in data:
            t = data["@type"]
            if isinstance(t, list): types.extend([str(x) for x in t])
            else: types.append(str(t))
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, dict) and "@type" in item:
                    t = item["@type"]
                    if isinstance(t, list): types.extend([str(x) for x in t])
                    else: types.append(str(t))
    # dedupe in order
    seen, out = set(), []
    for t in types:
        if t not in seen:
            seen.add(t); out.append(t)
    return out

def analyze_html_quality(html: str, url: str) -> Dict[str, Any]:
    soup = make_soup(html)

    # head/meta
    title = (soup.title.get_text(strip=True) if soup.title else "") or ""
    meta_desc_tag = soup.find("meta", attrs={"name": "description"})
    meta_desc = (meta_desc_tag.get("content") if meta_desc_tag else "") or ""
    meta_robots_tag = soup.find("meta", attrs={"name": "robots"})
    meta_robots = (meta_robots_tag.get("content") if meta_robots_tag else "") or ""
    link_canon = soup.find("link", rel=lambda v: v and "canonical" in (v if isinstance(v, list) else [v]))
    canonical = link_canon.get("href") if link_canon else ""

    # headings
    h1s = [h.get_text(strip=True) for h in soup.find_all("h1")]
    h2s = [h.get_text(strip=True) for h in soup.find_all("h2")]
    hns = []
    for n in range(1, 7):
        hns.extend([(n, h.get_text(strip=True)) for h in soup.find_all(f"h{n}")])

    jumps, last = [], None
    for lvl, txt in hns:
        if last is not None and (lvl - last) > 1:
            jumps.append((last, lvl, txt[:80]))
        last = lvl

    # content (stripped)
    stripped = make_soup(html)
    strip_boilerplate(stripped)
    text = stripped.get_text(separator="\n")
    words = _count_words(text)
    ratio = (len(text) / max(1, len(html)))

    # resources & schema
    links, pdfs = collect_links(html, url)
    # schema_types = _extract_ld_json_types(soup)

    issues, warnings, info = [], [], []
    if not title: issues.append("Missing <title>.")
    elif len(title) < 10 or len(title) > 75: warnings.append(f"Title length is {len(title)} chars (recommended ~10–70).")

    if not meta_desc: warnings.append("Missing meta description or empty content.")
    elif len(meta_desc) < 50 or len(meta_desc) > 170: warnings.append(f"Meta description length is {len(meta_desc)} chars (recommended ~50–160).")

    if meta_robots and any(t in meta_robots.lower() for t in ["noindex", "nofollow"]):
        warnings.append(f"robots meta contains: {meta_robots}")

    if not h1s: issues.append("Missing H1.")
    if len(h1s) > 1: warnings.append(f"Multiple H1s detected ({len(h1s)}).")
    if jumps: warnings.append("Heading level jumps found (e.g., h1 → h4).")
    if not soup.find("main"): warnings.append("No <main> landmark found.")
    if not canonical: warnings.append("Missing rel=canonical link.")

    # if not any("Product" in t for t in schema_types): warnings.append("No schema.org Product JSON-LD detected.")
    # else: info.append(f"Schema types: {', '.join(schema_types[:5])}")

    if words < 120: warnings.append(f"Low main-text word count (~{words}). The page may be JS-rendered or content-light.")
    if ratio < 0.05: warnings.append(f"Very low text-to-HTML ratio ({ratio:.2%}); likely heavy client-side rendering.")

    if len(pdfs) > 0: info.append(f"PDFs found: {len(pdfs)}")
    info.append(f"Total links: {len(links)}")

    return {
        "url": url, "title": title, "h1": (h1s[0] if h1s else ""),
        "word_count": words, "text_html_ratio": round(ratio, 4),
        "meta_description": meta_desc[:160], "robots": meta_robots,
        # "canonical": canonical, "schema_types": schema_types[:10],
        "issues": issues, "warnings": warnings, "info": info,
        "samples": {"first_h2": (h2s[0] if h2s else ""), "first_100_words": " ".join(text.split()[:100])},
    }




# ------------------------------ Prompt ---------------------------
# --- Role-aware prompts -------------------------------------------------------
BASE_SYSTEM_PROMPT = """You are an AI assistant producing crisp summaries of corporate product pages.
Be concise, factual, and business-focused. Ignore navigation, marketing fluff, and repetitive CTAs.
Do not list every link. If PDFs or external resources exist, mention them briefly as 'Additional resources available'.
Return Markdown with:
1) A 5–7 sentence Executive Summary (~280 words max)
2) 4–7 Key Highlights as bullets (value, differentiators, target users, notable metrics)
Avoid speculative claims. Be neutral and precise.
"""

# Audience lenses (adjust or extend as needed)
AUDIENCE_PRESETS = {
    "Executive (VP/Director)": {
        "label": "Executive (VP/Director)",
        "instructions": (
            "Write for VPs/Directors. Emphasize outcomes, ROI, risk mitigation, costs, "
            "time-to-value, customer impact, and strategic alignment. Avoid deep technical jargon."
        ),
    },
    "C‑suite (CEO/CFO/COO)": {
        "label": "C‑suite (CEO/CFO/COO)",
        "instructions": (
            "Write for CEOs/CFOs/COOs. Focus on business value, TCO, financial models (capex/opex), "
            "payback period, risk, and operational impact at the portfolio level."
        ),
    },
    "CTO / CIO": {
        "label": "CTO / CIO",
        "instructions": (
            "Write for CTOs/CIOs. Highlight architecture fit, integration complexity, interoperability, "
            "scalability, reliability/SLAs, and roadmap/modernization considerations."
        ),
    },
    "Senior Engineer / Architect": {
        "label": "Senior Engineer / Architect",
        "instructions": (
            "Write for senior engineers and architects. Be technically precise. Emphasize architecture, "
            "APIs/SDKs, performance characteristics, scale, deployment models, and reference architectures."
        ),
    },
    "Security / Compliance (CISO)": {
        "label": "Security / Compliance (CISO)",
        "instructions": (
            "Write for CISOs/security leads. Emphasize security posture, data handling, encryption, key management, "
            "identity/access controls, auditability, and compliance frameworks (e.g., ISO 27001, SOC 2)."
        ),
    },
    "Finance (CFO / FP&A)": {
        "label": "Finance (CFO / FP&A)",
        "instructions": (
            "Write for finance audiences. Emphasize TCO, licensing/pricing models, ROI, cost drivers, "
            "budget predictability, and financial risks."
        ),
    },
    "Product Management": {
        "label": "Product Management",
        "instructions": (
            "Write for PMs. Emphasize user value, JTBD, target segments/ICP, differentiation, proof points, "
            "KPIs/metrics, and roadmap implications."
        ),
    },
    "Sales / GTM": {
        "label": "Sales / GTM",
        "instructions": (
            "Write for sales leaders. Emphasize ICP, value propositions, competitive positioning, "
            "case studies, and common objections with crisp proof points."
        ),
    },
}

def build_system_prompt(audience_key: str, extra_instructions: str = "") -> str:
    preset = AUDIENCE_PRESETS.get(audience_key, {})
    lens = preset.get("instructions", "")
    system_prompt = BASE_SYSTEM_PROMPT
    if lens:
        system_prompt += "\n\nAudience lens:\n" + lens
    if extra_instructions.strip():
        system_prompt += "\n\nAdditional run-specific guidance:\n" + extra_instructions.strip()
    return system_prompt

USER_PROMPT_TEMPLATE = """\
URL: {url}
Title: {title}
H1/Product: {product}
Page Content (truncated):
\"\"\"
{content}
\"\"\"
Resources:
- PDFs detected: {pdf_count}
- Total links: {link_count}
Instructions:
- Summarize the product's purpose, business value, key capabilities/differentiators, and ideal audience.
- Do NOT include navigation/menu text.
- Do NOT list all links; simply note that additional resources are available if present.
- Keep the language concise and executive-friendly.
"""

def build_messages(
    url: str,
    title: Optional[str],
    product: Optional[str],
    content: str,
    pdf_count: int,
    link_count: int,
    system_prompt: str,   # <-- new
) -> List[Dict]:
    prompt = USER_PROMPT_TEMPLATE.format(
        url=url,
        title=title or "N/A",
        product=product or "N/A",
        content=content,
        pdf_count=pdf_count,
        link_count=link_count,
    )
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": prompt},
    ]


def split_output(markdown: str):
    """
    Split the LLM markdown into:
      - exec_part: the first paragraph/section (keeps line breaks & blank lines)
      - bullets: up to 7 bullet points extracted from the rest
    """
    md = (markdown or "").replace("\r\n", "\n").replace("\r", "\n").strip()

    # Collect bullets (single level)
    bullets: List[str] = []
    for line in md.splitlines():
        if re.match(r"^\s*[-\*\•]\s+", line):
            bullets.append(re.sub(r"^\s*[-\*\•]\s+", "", line).rstrip())

    # Take first paragraph after an optional top-level heading
    parts = md.split("\n\n")
    if parts and re.match(r"^\s*#{1,6}\s", parts[0]):
        parts = parts[1:]

    exec_part = parts[0] if parts else md

    # Intra-line cleanup (keep newlines, remove only extra spaces/tabs inside lines)
    exec_part = "\n".join(re.sub(r"[ \t]+", " ", ln).rstrip() for ln in exec_part.split("\n"))

    # Collapse excessive blank lines but preserve paragraph breaks
    exec_part = re.sub(r"\n{3,}", "\n\n", exec_part).strip()

    # Keep at most 7 bullets


# ------------------------- Domain Allowlist ----------------------
def is_domain_allowed(url: str, allowed: List[str]) -> bool:
    if not allowed:
        return True
    try:
        host = urlparse(url).hostname or ""
    except Exception:
        return False
    host = host.lower()
    for dom in allowed:
        d = dom.strip().lower()
        if not d:
            continue
        if host == d or host.endswith("." + d):
            return True
    return False

# --------------------------- LLM Plumbing ------------------------
def build_headers(auth_type: str, api_key: Optional[str], extra_headers_str: str) -> Dict[str, str]:
    headers: Dict[str, str] = {"Content-Type": "application/json"}
    if auth_type == "Authorization (Bearer)":
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
    elif auth_type in ("api-key", "x-api-key"):
        if api_key:
            headers[auth_type] = api_key
    # Merge custom headers JSON
    if extra_headers_str.strip():
        try:
            headers.update(json.loads(extra_headers_str))
        except Exception as e:
            raise RuntimeError(f"Extra headers JSON invalid: {e}")
    return headers

def messages_to_prompt(messages: List[Dict]) -> str:
    sys_parts = [m["content"] for m in messages if m.get("role") == "system"]
    user_parts = [m["content"] for m in messages if m.get("role") == "user"]
    prompt = ""
    if sys_parts:
        prompt += f"System:\n{sys_parts[-1]}\n\n"
    prompt += "\n\n".join([f"User:\n{u}" for u in user_parts]) + "\nAssistant:\n"
    return prompt

def call_llm_requests(
    base_url: str,
    model: str,
    messages: List[Dict],
    headers: Dict[str, str],
    timeout_seconds: int,
    verify_param,  # bool or str path for CA bundle
    max_output_tokens: int = 700,
) -> str:
    if not base_url.endswith("/v1"):
        raise RuntimeError("Base URL missing")
    chat_url = base_url.rstrip("/") + "/chat/completions"
    comp_url = base_url.rstrip("/") + "/completions"

    # 1) Try /chat/completions
    chat_payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": max_output_tokens,
    }
    try:
        r = requests.post(chat_url, headers=headers, json=chat_payload,
                          timeout=timeout_seconds, verify=verify_param)
        if r.status_code == 200:
            data = r.json()
            # standard OpenAI-style
            if data.get("choices") and "message" in data["choices"][0]:
                return (data["choices"][0]["message"]["content"] or "").strip()
            # Some servers return "text" in chat (rare)
            if data.get("choices") and "text" in data["choices"][0]:
                return (data["choices"][0]["text"] or "").strip()
            raise RuntimeError(f"Chat response missing expected fields: {str(data)[:300]}")
        elif r.status_code in (404, 405):
            # not implemented, fall back to /completions
            pass
        elif r.status_code == 401:
            raise RuntimeError("401 Unauthorized from /chat/completions. Check auth header type and API key.")
        else:
            raise RuntimeError(f"/chat/completions -> {r.status_code}: {r.text[:400]}")
    except requests.exceptions.RequestException as e:
        # Network/TLS/timeout issues
        raise RuntimeError(f"/chat/completions request failure: {repr(e)}")

    # 2) Fallback: /completions
    prompt = messages_to_prompt(messages)
    comp_payload = {
        "model": model,
        "prompt": prompt,
        "temperature": 0.2,
        "max_tokens": max_output_tokens,
    }
    try:
        r2 = requests.post(comp_url, headers=headers, json=comp_payload,
                           timeout=timeout_seconds, verify=verify_param)
        if r2.status_code == 200:
            data2 = r2.json()
            if data2.get("choices") and "text" in data2["choices"][0]:
                return (data2["choices"][0]["text"] or "").strip()
            raise RuntimeError(f"Completions response missing expected fields: {str(data2)[:300]}")
        elif r2.status_code == 401:
            raise RuntimeError("401 Unauthorized from /completions. Check auth header type and API key.")
        else:
            raise RuntimeError(f"/completions -> {r2.status_code}: {r2.text[:400]}")
    except requests.exceptions.RequestException as e:
        raise RuntimeError(f"/completions request failure: {repr(e)}")

def call_llm(
    base_url: str,
    model: str,
    messages: List[Dict],
    headers: Dict[str, str],
    timeout_seconds: int,
    verify_param,
    max_output_tokens: int = 700,
) -> str:
    """
    Prefer your llm_client if available, except when a custom CA bundle path is provided,
    in which case we call requests directly to pass `verify=<path>`.
    """
    # If verify is a PEM path, use the requests-based path (llm_client only supports boolean verify)
    if isinstance(verify_param, str) and verify_param.strip():
        return call_llm_requests(
            base_url=base_url,
            model=model,
            messages=messages,
            headers=headers,
            timeout_seconds=timeout_seconds,
            verify_param=verify_param,
            max_output_tokens=max_output_tokens,
        )


try:
    from playwright.sync_api import sync_playwright
    _HAS_PLAYWRIGHT = True
except Exception:
    sync_playwright = None
    _HAS_PLAYWRIGHT = False


def fetch_url_rendered(url: str, timeout: int = 25, verify=True) -> Tuple[str, str]:
    # NOTE: 'verify' is not used here; Playwright manages TLS internally.
    if not _HAS_PLAYWRIGHT:
        
        raise RuntimeError(
            "Playwright is not installed. To enable JS rendering install it with: \n"
            "pip install playwright\npython -m playwright install chromium"
        )

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            ctx = browser.new_context(ignore_https_errors=(verify is False))
            page = ctx.new_page()
            page.set_default_timeout(timeout * 1000)
            page.goto(url, wait_until="networkidle")
            final = page.url
            html = page.content()
            return final, html
        finally:
            browser.close()

def test_connectivity(base_url: str, headers: Dict[str, str], verify_param, model_for_test: str) -> Tuple[bool, str]:
    if not base_url.endswith("/v1"):
        return False, "Base URL must include /v1 (e.g., https://host/v1)"
    try:
        # GET /models (often open without auth, but still useful)
        r = requests.get(base_url.rstrip("/") + "/models", headers=headers, timeout=15, verify=verify_param)
        if r.status_code >= 400:
            return False, f"/models -> {r.status_code}: {r.text[:200]}"
        # POST /chat/completions (detects 401 or 404 quickly)
        chat_payload = {
            "model": model_for_test or "dummy",
            "messages": [{"role":"user", "content":"ping"}],
            "max_tokens": 5,
            "temperature": 0.0
        }
        r2 = requests.post(base_url.rstrip("/") + "/chat/completions", headers=headers,
                           json=chat_payload, timeout=15, verify=verify_param)
        if r2.status_code == 200:
            return True, "Connectivity OK (chat/completions)."
        elif r2.status_code == 401:
            return False, "401 Unauthorized on /chat/completions. Check auth header type and API key."
        elif r2.status_code in (404, 405):
            # Try /completions
            comp_payload = {
                "model": model_for_test or "dummy",
                "prompt": "ping",
                "max_tokens": 5,
                "temperature": 0.0
            }
            r3 = requests.post(base_url.rstrip("/") + "/completions", headers=headers,
                               json=comp_payload, timeout=15, verify=verify_param)
            if r3.status_code == 200:
                return True, "Connectivity OK via /completions (fallback)."
            elif r3.status_code == 401:
                return False, "401 Unauthorized on /completions. Check auth header type and API key."
            return False, f"/completions -> {r3.status_code}: {r3.text[:200]}"
        return False, f"/chat/completions -> {r2.status_code}: {r2.text[:200]}"
    except requests.exceptions.SSLError as e:
        return False, f"SSL error: {e}. Try providing a CA bundle or enable 'Skip TLS verification' for testing."
    except requests.exceptions.RequestException as e:
        return False, f"Connection error: {repr(e)}"


import json as _json
from typing import List, Tuple, Optional

def split_output_safe(markdown: Optional[str]) -> Tuple[str, List[str]]:
    """
    Safe wrapper around split_output(); always returns (exec, bullets).
    Falls back to a basic parser if split_output raises or returns invalid values.
    """
    try:
        if markdown is None:
            raise ValueError("markdown=None")
        exec_part, bullets = split_output(markdown)  # your existing function
        if exec_part is None or bullets is None:
            raise ValueError("split_output returned None")
        if not isinstance(bullets, list):
            raise ValueError("split_output returned non-list bullets")
        return exec_part, bullets[:7]
    except Exception:
        # Fallback: naive extraction
        md = (markdown or "")
        # try to get bullets from common patterns
        bullets = []
        for line in md.splitlines():
            line = line.rstrip()
            if line.strip().startswith(("-", "*", "•")):
                bullets.append(line.lstrip("-*• ").strip())
        bullets = [b for b in bullets if b][:7]

        # take the first paragraph as summary
        parts = md.strip().split("\n\n")
        exec_part = parts[0].strip() if parts else md.strip()
        return exec_part, bullets


def parse_llm_output(llm_out: Optional[str]) -> Tuple[str, List[str]]:
    """
    Accepts either Markdown or JSON shaped output from the model.
    Always returns (exec_summary, key_bullets).
    """
    if not isinstance(llm_out, str):
        raise RuntimeError(f"LLM returned non-text payload: {type(llm_out).__name__}")

    text = llm_out.strip()
    if not text:
        # Empty response: return empty exec + no bullets; UI can handle this gracefully.
        return "", []

    # Try JSON first (many models sometimes switch to JSON)
    lead = text.lstrip()
    if lead.startswith("{") or lead.startswith("["):
        try:
            data = _json.loads(lead)
            if isinstance(data, dict):
                exec_summary = (data.get("executive_summary")
                                or data.get("summary")
                                or data.get("Executive Summary")
                                or "")
                bullets = (data.get("key_highlights")
                           or data.get("bullets")
                           or data.get("Key Highlights")
                           or [])
                if isinstance(bullets, str):
                    bullets = [b.strip(" -*•") for b in bullets.splitlines() if b.strip()]
                elif not isinstance(bullets, list):
                    bullets = []
                return exec_summary, bullets[:7]
            if isinstance(data, list):
                # Assume list of bullets with no exec paragraph
                return "", [str(x) for x in data][:7]
        except Exception:
            # fall through to markdown parse
            pass

    # Parse as Markdown
    return split_output_safe(text)


def summarize_url(
    url: str,
    model: str,
    base_url: str,
    headers: Dict[str, str],
    verify_param,
    timeout_seconds: int,
    system_prompt: str,
) -> ExecSummary:
    # Fetch
    final_url, html = _ensure_pair(fetch_url(url, timeout=20, verify=verify_param), "fetch_url")

    # Title/H1
    title, h1 = _ensure_pair(get_title_and_h1(html), "get_title_and_h1")
    product = infer_product_name(title, h1)

    # Main text (always a string)
    try:
        main_text = readability_extract(html)
        if len((main_text or "").split()) < 50:
            main_text = simple_extract(html)
    except Exception:
        main_text = simple_extract(html)

    # Links/PDFs
    links, pdfs = _ensure_pair(collect_links(html, final_url), "collect_links")

    content = truncate_for_model(main_text or "", max_chars=12000)

    # (unchanged) build messages → call LLM → split output → return ExecSummary
    msgs = build_messages(
        final_url, title, product, content,
        pdf_count=len(pdfs), link_count=len(links),
        system_prompt=system_prompt
    )

    llm_out = call_gemini(
        base_url=base_url,
        model=model,
        messages=msgs,
        headers=headers,
        timeout_seconds=timeout_seconds,
        verify_param=verify_param,
        max_output_tokens=4096,
    )

    if not isinstance(llm_out, str):
        raise RuntimeError(f"LLM returned non-text: {type(llm_out).__name__}")

    exec_summary, key_bullets = parse_llm_output(llm_out)

    resources_note = None
    samples: List[str] = []
    if len(pdfs) > 0 or len(links) > 0:
        resources_note = "Additional resources available on the page (PDFs, reference links)."
        samples = (pdfs[:2] if pdfs else []) + [l for l in links if not l.lower().endswith(".pdf")][:2]

    return ExecSummary(
        url=final_url,
        title=title,
        product=product,
        executive_summary=exec_summary,
        key_highlights=key_bullets,
        resources_note=resources_note,
        sample_resources=samples,
    )

# --------------------------- Export helpers -----------------------
def build_docx(summaries: List[ExecSummary]) -> bytes:
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, s in enumerate(summaries):
        if i > 0:
            doc.add_page_break()
        doc.add_heading(s.product or s.title or s.url, level=1)

        p = doc.add_paragraph()
        r = p.add_run("URL: ")
        r.bold = True
        p.add_run(s.url)

        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(s.executive_summary)

        if s.key_highlights:
            doc.add_heading("Key Highlights", level=2)
            for b in s.key_highlights:
                doc.add_paragraph(b).style = doc.styles['List Bullet']

        if s.resources_note:
            doc.add_paragraph()
            rp = doc.add_paragraph()
            run = rp.add_run("Note: ")
            run.bold = True
            rp.add_run(s.resources_note)

        if s.sample_resources:
            doc.add_paragraph("Sample resources:")
            for link in s.sample_resources:
                doc.add_paragraph(link).style = doc.styles['List Bullet']

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ------------------------------ UI --------------------------------
st.set_page_config(page_title="Executive Summarizer", page_icon="📝", layout="wide")
ip = get_user_ip()
allowed, remaining, wait_seconds = check_and_update_access(ip)
if not allowed:
    st.error(f"You have reached the limit of {ACCESS_LIMIT} accesses in 24 hours. Please try again in {wait_seconds//3600} hours.")
    st.stop()
else:
    st.info(f"Notice: You are limited to {ACCESS_LIMIT} accesses per 24 hours (based on your IP address). Remaining: {remaining}")
st.title("Executive Summarizer for Any Web page")
st.title("Executive Summarizer for Any Web page")
base_url=os.getenv("OPENAI_BASE_URL", "")
# model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
api_key = os.getenv("OPENAI_API_KEY", "")
auth_type ="Authorization (Bearer)"
extra_headers_str=""
skip_verify_toggle=os.getenv("OPENAI_SKIP_VERIFY", "false").lower() == "true"
ca_bundle_path=""
timeout_seconds = 60
allowed_domains_env = os.getenv("ALLOWED_DOMAINS", "")
allowed_domains_in = ""
allowed_domains = [d.strip() for d in allowed_domains_in.split(",") if d.strip()]
verify_param = False if skip_verify_toggle else (ca_bundle_path if ca_bundle_path.strip() else True)

# Build headers
try:
    hdrs = build_headers(auth_type, api_key if auth_type != "None" else None, extra_headers_str)
except Exception as e:
    hdrs = None
    st.error(str(e))


with st.sidebar:
    model = st.text_input("Model", value=os.getenv("OPENAI_MODEL", "gemini-2.5-flash"))

    # --- New: Audience role & extra guidance ---
    audience = st.selectbox(
        "Audience",
        options=list(AUDIENCE_PRESETS.keys()),
        index=0,  # default to "Executive (VP/Director)"
        format_func=lambda k: AUDIENCE_PRESETS[k]["label"]
    )
    extra_guidance = st.text_area(
        "Optional: add run-specific guidance",
        placeholder="e.g., compare briefly with alternatives; highlight specifics"
    )

    # Build the final system prompt for this run
    system_prompt = build_system_prompt(audience, extra_instructions=extra_guidance)

    # Optional: debugging aid
    with st.expander("Show effective system prompt", expanded=False):
        st.code(system_prompt, language="markdown")

    # Playwright availability note
    try:
        _has_pw = _HAS_PLAYWRIGHT  # set earlier near fetch_url_rendered
    except NameError:
        _has_pw = False
    if not _has_pw:
        st.info(
            "JS rendering (Playwright) is not available in this environment.\n"
            "To enable it, install playwright: `pip install playwright` and run `python -m playwright install chromium`."
        )

# , "📦 Batch (CSV)", "🛠️ Gap analysis (dev)"]
tabs = st.tabs(["🔗 Single URL"])

# ---- Single URL Tab ----
with tabs[0]:
    st.subheader("Summarize a single page")
    url = st.text_input("Page URL", placeholder="https://www.website.com/...")
    run_btn = st.button("Summarize", type="primary")

    if run_btn:
            if not url or not (url.startswith("http://") or url.startswith("https://")):
                st.error("Please enter a valid URL, including the 'http://' or 'https://' prefix.")
                st.stop()

            res = None
            with st.status("Fetching and summarizing...", expanded=True) as status:
                # 1) Upstream steps (fetch + extraction)
                try:
                    res = summarize_url(
                        url=url,
                        model=model,
                        base_url=base_url,
                        headers=hdrs,
                        verify_param=verify_param,
                        timeout_seconds=int(timeout_seconds),
                        system_prompt=system_prompt,
                    )
                    status.update(label="Done", state="complete")
                except RetryError as e:
                    st.error("Failed fetching the page (before LLM). Check TLS / proxy / allowlist settings.")
                    st.exception(e)
                    res = None
                except RuntimeError as e:
                    # This will catch _ensure_pair errors and other pre-LLM pipeline issues.
                    st.error(f"Pre‑LLM pipeline failure: {e}")
                    res = None
                except Exception as e:
                    # If it fails during the LLM call, your summarize_url can raise here;
                    # otherwise keep your existing message.
                    st.error(f"Failed calling LLM. Base URL: {base_url}, Model: {model}")
                    st.exception(e)
                    res = None
            if res is not None:
                st.markdown(f"### {res.product or res.title or 'Executive Summary'}")
                st.markdown(f"**URL:** {res.url}")

                # st.markdown("#### Executive Summary")
                st.markdown(res.executive_summary or "_(no summary returned)_")

                if res.key_highlights:
                    st.markdown("#### Key Highlights")
                    for b in res.key_highlights:
                        st.markdown(f"- {b}")

                if res.resources_note:
                    st.info(res.resources_note)

                if res.sample_resources:
                    with st.expander("Sample resources"):
                        for r in res.sample_resources:
                            st.markdown(f"- {r}")

                # Downloads
                col1, col2, col3 = st.columns(3)
                with col1:
                    md = (
                        f"# {res.product or res.title or 'Executive Summary'}\n\n"
                        f"**URL:** {res.url}\n\n"
                        f"## Executive Summary\n\n{res.executive_summary}\n\n"
                        f"## Key Highlights\n" + "\n".join([f"- {b}" for b in res.key_highlights])
                    )
                    st.download_button("⬇️ Download Markdown", data=md.encode("utf-8"),
                                       file_name="summary.md", mime="text/markdown")
                with col2:
                    # `res` is already an ExecSummary instance
                    st.download_button("⬇️ Download JSON",
                        data=res.model_dump_json(indent=2).encode("utf-8"),
                        file_name="summary.json", mime="application/json")
                with col3:
                    doc_bytes = build_docx([res])
                    st.download_button("⬇️ Download DOCX", data=doc_bytes,
                        file_name="summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
